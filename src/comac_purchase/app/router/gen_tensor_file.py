"""
LLM 工具相关的 API 路由
"""
import asyncio
import uuid
import logging
from pathlib import Path
from typing import Annotated, Dict, Any
import sys

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel
from sqlalchemy.orm import Session

# 添加项目根目录到路径
project_root = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(project_root))

from ...config import settings
from ...db import File, TenderGeneration, Project, get_db
from ...model_session.model_session_manager import get_manager, SessionStatus

# 配置日志
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
if not logger.handlers:
    handler = logging.StreamHandler()
    handler.setLevel(logging.INFO)
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    handler.setFormatter(formatter)
    logger.addHandler(handler)

router = APIRouter(prefix="/llm-tool", tags=["llm-tool"])


class GenerateTenderRequest(BaseModel):
    """生成采购征询文件请求模型"""
    project_id: int  # 项目ID（必需）
    technical_requirement_file_id: str  # 业务需求文件ID
    procurement_requirement_file_id: str  # 采购部门要求文件ID


def read_docx_to_text(file_path: Path) -> str:
    """读取 docx 文件内容并转换为纯文本（保留结构）"""
    from docx import Document
    
    doc = Document(file_path)
    text_lines = []
    
    for element in doc.element.body:
        # 处理段落
        if element.tag.endswith('p'):
            from docx.oxml.text.paragraph import CT_P
            from docx.text.paragraph import Paragraph
            
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                text = para.text.strip()
                if text:
                    # 检查是否是标题
                    if para.style and ('Heading' in para.style.name or '标题' in para.style.name):
                        level = _get_heading_level(para.style.name)
                        if level > 0:
                            text_lines.append('#' * level + ' ' + text)
                        else:
                            text_lines.append(text)
                    else:
                        text_lines.append(text)
        
        # 处理表格
        elif element.tag.endswith('tbl'):
            from docx.oxml.table import CT_Tbl
            from docx.table import Table
            
            if isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table_text = _table_to_text(table)
                if table_text:
                    text_lines.append('')
                    text_lines.append(table_text)
                    text_lines.append('')
    
    return '\n'.join(text_lines)


def _get_heading_level(style_name: str) -> int:
    """从样式名称获取标题级别"""
    if 'Heading 1' in style_name or '标题 1' in style_name or 'Title' in style_name:
        return 1
    elif 'Heading 2' in style_name or '标题 2' in style_name:
        return 2
    elif 'Heading 3' in style_name or '标题 3' in style_name:
        return 3
    elif 'Heading 4' in style_name or '标题 4' in style_name:
        return 4
    elif 'Heading 5' in style_name or '标题 5' in style_name:
        return 5
    elif 'Heading 6' in style_name or '标题 6' in style_name:
        return 6
    return 0


def _table_to_text(table) -> str:
    """将表格转换为文本格式"""
    text_lines = []
    for row in table.rows:
        row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        text_lines.append(' | '.join(row_data))
    return '\n'.join(text_lines)


def read_markdown_file(file_path: Path) -> str:
    """读取 markdown 文件内容"""
    return file_path.read_text(encoding='utf-8')


def markdown_to_docx(markdown_content: str, output_path: Path):
    """将 Markdown 内容转换为 docx 文件"""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 设置标题样式
    for i in range(1, 7):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = '黑体'
        heading_font.size = Pt(18 - i * 2)
        heading_font.bold = True
    
    lines = markdown_content.split('\n')
    i = 0
    
    def add_formatted_text(paragraph, text: str):
        """向段落添加格式化的文本（支持粗体）"""
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part:
                paragraph.add_run(part)
    
    while i < len(lines):
        line = lines[i]
        stripped_line = line.strip()
        
        if not stripped_line:
            i += 1
            continue
        
        # 处理标题
        if stripped_line.startswith('#'):
            level = 0
            while level < len(stripped_line) and stripped_line[level] == '#':
                level += 1
            title_text = stripped_line[level:].strip()
            title_text = re.sub(r'\*\*(.*?)\*\*', r'\1', title_text)
            
            if level == 1:
                heading = doc.add_heading(title_text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(title_text, min(level - 1, 6))
        
        # 处理表格
        elif stripped_line.startswith('|') and '|' in stripped_line[1:]:
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if not re.match(r'^\|[\s\-:]+\|', row_line):
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    cells = [re.sub(r'\*\*(.*?)\*\*', r'\1', cell) for cell in cells]
                    table_rows.append(cells)
                i += 1
            i -= 1
            
            if table_rows:
                num_cols = max(len(row) for row in table_rows) if table_rows else 1
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                if table_rows:
                    header_cells = table.rows[0].cells
                    for j in range(num_cols):
                        if j < len(table_rows[0]):
                            header_cells[j].text = table_rows[0][j]
                        else:
                            header_cells[j].text = ''
                    
                    for row_data in table_rows[1:]:
                        row = table.add_row()
                        for j in range(num_cols):
                            if j < len(row_data):
                                row.cells[j].text = row_data[j]
                            else:
                                row.cells[j].text = ''
        
        # 处理无序列表
        elif re.match(r'^[\-\*\+]\s+', stripped_line):
            list_text = re.sub(r'^[\-\*\+]\s+', '', stripped_line)
            p = doc.add_paragraph(list_text, style='List Bullet')
            if '**' in list_text:
                p.clear()
                add_formatted_text(p, list_text)
        
        # 处理有序列表
        elif re.match(r'^\d+\.\s+', stripped_line):
            list_text = re.sub(r'^\d+\.\s+', '', stripped_line)
            p = doc.add_paragraph(list_text, style='List Number')
            if '**' in list_text:
                p.clear()
                add_formatted_text(p, list_text)
        
        # 处理普通段落
        else:
            if line.startswith('   ') or line.startswith('\t'):
                text = stripped_line
                p = doc.add_paragraph(text)
                if '**' in text:
                    p.clear()
                    add_formatted_text(p, text)
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, stripped_line)
        
        i += 1
    
    # 保存文件
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


async def _generate_tender_task(
    tender_generation_id: int,
    technical_file_id: str,
    procurement_file_id: str,
    session_id: str
):
    """后台任务：生成采购征询文件"""
    from ...db import SessionLocal
    
    logger.info(f"[生成任务 {tender_generation_id}] 开始执行生成任务，会话ID: {session_id}")
    logger.info(f"[生成任务 {tender_generation_id}] 技术文件ID: {technical_file_id}, 采购文件ID: {procurement_file_id}")
    
    # 创建新的数据库会话
    db = SessionLocal()
    try:
        # 获取生成采购征询文件记录
        tender_generation = db.query(TenderGeneration).filter(TenderGeneration.id == tender_generation_id).first()
        if not tender_generation:
            logger.error(f"[生成任务 {tender_generation_id}] 生成记录不存在")
            raise RuntimeError(f"生成采购征询文件记录不存在: {tender_generation_id}")
        
        logger.info(f"[生成任务 {tender_generation_id}] 获取生成记录成功，当前状态: {tender_generation.status}")
        
        # 确保状态为 running（创建时已经是 running，这里只是确认）
        if tender_generation.status != 'running':
            logger.info(f"[生成任务 {tender_generation_id}] 更新状态为 running")
            tender_generation.status = 'running'
            db.commit()
        
        # 重新查询文件（使用新的数据库会话）
        technical_file = db.query(File).filter(File.file_id == technical_file_id).first()
        procurement_file = db.query(File).filter(File.file_id == procurement_file_id).first()
        
        if not technical_file or not procurement_file:
            logger.error(f"[生成任务 {tender_generation_id}] 文件记录不存在 - 技术文件: {technical_file is not None}, 采购文件: {procurement_file is not None}")
            tender_generation.status = 'failed'
            db.commit()
            raise RuntimeError("文件不存在")
        
        logger.info(f"[生成任务 {tender_generation_id}] 文件记录查询成功 - 技术文件: {technical_file.origin_name}, 采购文件: {procurement_file.origin_name}")
        
        # 读取文件内容
        files_folder = Path(settings.data_folder) / "files"
        technical_file_path = files_folder / technical_file.file_name
        procurement_file_path = files_folder / procurement_file.file_name
        
        logger.info(f"[生成任务 {tender_generation_id}] 文件路径 - 技术文件: {technical_file_path}, 采购文件: {procurement_file_path}")
        
        if not technical_file_path.exists() or not procurement_file_path.exists():
            logger.error(f"[生成任务 {tender_generation_id}] 物理文件不存在 - 技术文件存在: {technical_file_path.exists()}, 采购文件存在: {procurement_file_path.exists()}")
            tender_generation.status = 'failed'
            db.commit()
            raise FileNotFoundError("文件不存在")
        
        logger.info(f"[生成任务 {tender_generation_id}] 物理文件验证成功")
        
        # 使用 model_session 生成内容
        manager = get_manager()
        
        # 等待会话启动（最多等待20分钟）
        logger.info(f"[生成任务 {tender_generation_id}] 等待会话启动: {session_id}")
        wait_count = 0
        max_wait_start = 1200  # 最多等待20分钟让会话启动
        while wait_count < max_wait_start:
            try:
                session_status = manager.get_session_status(session_id)
                if session_status in [SessionStatus.RUNNING, SessionStatus.FINISHED, SessionStatus.ERROR]:
                    logger.info(f"[生成任务 {tender_generation_id}] 会话已启动，状态: {session_status}")
                    break
            except Exception as e:
                # 会话可能还未创建，继续等待
                pass
            await asyncio.sleep(1)
            wait_count += 1
        
        if wait_count >= max_wait_start:
            # 检查会话是否存在
            logger.warning(f"[生成任务 {tender_generation_id}] 会话启动等待超时（{max_wait_start}秒），检查会话状态")
            session = manager.get_session(session_id)
            if not session:
                logger.error(f"[生成任务 {tender_generation_id}] 会话不存在: {session_id}")
                tender_generation.status = 'failed'
                db.commit()
                raise RuntimeError(f"会话启动超时，会话不存在: {session_id}")
            else:
                status = manager.get_session_status(session_id)
                logger.error(f"[生成任务 {tender_generation_id}] 会话启动超时，当前状态: {status}")
                tender_generation.status = 'failed'
                db.commit()
                raise RuntimeError(f"会话启动超时，当前状态: {status}")
        
        # 等待会话完成，添加超时机制（最多等待20分钟）
        logger.info(f"[生成任务 {tender_generation_id}] 等待会话完成: {session_id}")
        max_wait_time = 1200  # 20分钟
        wait_count = 0
        last_status = None
        
        while wait_count < max_wait_time:
            try:
                session_status = manager.get_session_status(session_id)
                
                # 记录状态变化
                if session_status != last_status:
                    logger.info(f"[生成任务 {tender_generation_id}] 会话状态变化: {last_status} -> {session_status}")
                    last_status = session_status
                
                if session_status == SessionStatus.FINISHED:
                    logger.info(f"[生成任务 {tender_generation_id}] 会话完成: {session_id}")
                    # 等待一小段时间，确保文件已保存到磁盘
                    await asyncio.sleep(2)
                    break
                elif session_status == SessionStatus.ERROR:
                    logger.error(f"[生成任务 {tender_generation_id}] 会话执行失败，状态: ERROR")
                    tender_generation.status = 'failed'
                    db.commit()
                    # 尝试获取错误信息
                    session = manager.get_session(session_id)
                    error_msg = "会话执行失败"
                    if session and hasattr(session, 'error_message'):
                        error_msg = f"会话执行失败: {session.error_message}"
                        logger.error(f"[生成任务 {tender_generation_id}] 错误信息: {session.error_message}")
                    raise RuntimeError(error_msg)
                elif session_status == SessionStatus.RUNNING:
                    # 状态是 running，继续等待
                    await asyncio.sleep(1)
                    wait_count += 1
                    continue
                else:
                    # 其他状态（如 IDLE, STOPPED），等待一下再检查
                    await asyncio.sleep(1)
                    wait_count += 1
            except Exception as e:
                # 如果获取状态失败，检查会话是否还存在
                logger.warning(f"[生成任务 {tender_generation_id}] 获取会话状态时出错: {str(e)}")
                session = manager.get_session(session_id)
                if not session:
                    logger.error(f"[生成任务 {tender_generation_id}] 会话不存在: {session_id}")
                    tender_generation.status = 'failed'
                    db.commit()
                    raise RuntimeError(f"会话不存在: {session_id}, 错误: {str(e)}")
                # 继续等待
                await asyncio.sleep(1)
                wait_count += 1
        
        if wait_count >= max_wait_time:
            logger.error(f"[生成任务 {tender_generation_id}] 等待会话完成超时（{max_wait_time}秒），当前状态: {last_status}")
            tender_generation.status = 'failed'
            db.commit()
            raise RuntimeError(f"等待会话完成超时（{max_wait_time}秒），当前状态: {last_status}")
        
        # 获取会话并提取内容
        logger.info(f"[生成任务 {tender_generation_id}] 开始提取生成内容")
        session = manager.get_session(session_id)
        if not session:
            logger.error(f"[生成任务 {tender_generation_id}] 会话不存在: {session_id}")
            tender_generation.status = 'failed'
            db.commit()
            raise RuntimeError(f"会话不存在: {session_id}")
        
        # 记录会话类型和状态
        session_type = type(session).__name__
        logger.info(f"[生成任务 {tender_generation_id}] 会话类型: {session_type}, 状态: {session.status}")
        
        # 从会话中提取完整内容
        # 支持 LiveSession（从 response_list）和 HistorySession（从 content）
        markdown_content = ""
        
        # 方法1：如果是 LiveSession，尝试从 response_list 获取
        if hasattr(session, 'response_list'):
            if session.response_list:
                chunk_count = 0
                for chunk in session.response_list:
                    if chunk and hasattr(chunk, 'delta_content') and chunk.delta_content:
                        markdown_content += chunk.delta_content
                        chunk_count += 1
                logger.info(f"[生成任务 {tender_generation_id}] 从 response_list 提取内容完成，共 {chunk_count} 个chunk，内容长度: {len(markdown_content)} 字符")
            else:
                logger.warning(f"[生成任务 {tender_generation_id}] LiveSession 的 response_list 为空")
        
        # 方法2：如果是 HistorySession，从 content 属性获取
        if not markdown_content and hasattr(session, 'content'):
            if session.content:
                markdown_content = session.content
                logger.info(f"[生成任务 {tender_generation_id}] 从 content 属性提取内容完成，内容长度: {len(markdown_content)} 字符")
            else:
                logger.warning(f"[生成任务 {tender_generation_id}] HistorySession 的 content 为空，尝试从文件重新加载")
                # 如果HistorySession的content为空，尝试从文件重新加载
                if session_type == 'HistorySession':
                    try:
                        from ...model_session.model_session_manager import HistorySession
                        reloaded_session = HistorySession.from_file(session_id)
                        if reloaded_session and reloaded_session.content:
                            markdown_content = reloaded_session.content
                            logger.info(f"[生成任务 {tender_generation_id}] 从文件重新加载内容成功，内容长度: {len(markdown_content)} 字符")
                    except Exception as e:
                        logger.warning(f"[生成任务 {tender_generation_id}] 从文件重新加载失败: {str(e)}")
        
        # 方法3：如果前两种方法都失败，使用 get_response() 方法获取
        if not markdown_content:
            try:
                logger.info(f"[生成任务 {tender_generation_id}] 尝试使用 get_response() 方法提取内容")
                chunk_count = 0
                async for chunk in session.get_response():
                    if chunk and hasattr(chunk, 'delta_content') and chunk.delta_content:
                        markdown_content += chunk.delta_content
                        chunk_count += 1
                logger.info(f"[生成任务 {tender_generation_id}] 从 get_response() 提取内容完成，共 {chunk_count} 个chunk，内容长度: {len(markdown_content)} 字符")
            except Exception as e:
                logger.warning(f"[生成任务 {tender_generation_id}] 使用 get_response() 提取内容失败: {str(e)}")
        
        # 如果仍然没有内容，记录详细的调试信息
        if not markdown_content:
            logger.error(f"[生成任务 {tender_generation_id}] 未能获取生成的内容（已尝试所有方法）")
            logger.error(f"[生成任务 {tender_generation_id}] 会话类型: {session_type}")
            logger.error(f"[生成任务 {tender_generation_id}] 会话状态: {session.status}")
            logger.error(f"[生成任务 {tender_generation_id}] 是否有 response_list: {hasattr(session, 'response_list')}")
            if hasattr(session, 'response_list'):
                logger.error(f"[生成任务 {tender_generation_id}] response_list 长度: {len(session.response_list) if session.response_list else 0}")
            logger.error(f"[生成任务 {tender_generation_id}] 是否有 content: {hasattr(session, 'content')}")
            if hasattr(session, 'content'):
                logger.error(f"[生成任务 {tender_generation_id}] content 是否为空: {not session.content if session.content else True}")
            tender_generation.status = 'failed'
            db.commit()
            raise RuntimeError("未能获取生成的内容")
        
        # 转换为 docx
        logger.info(f"[生成任务 {tender_generation_id}] 开始转换为docx格式")
        files_folder = Path(settings.data_folder) / "files"
        file_id = str(uuid.uuid4())
        file_name = str(uuid.uuid4())
        output_path = files_folder / file_name
        
        try:
            markdown_to_docx(markdown_content, output_path)
            logger.info(f"[生成任务 {tender_generation_id}] docx文件生成成功: {output_path}")
        except Exception as e:
            logger.error(f"[生成任务 {tender_generation_id}] docx转换失败: {str(e)}")
            raise
        
        # 创建文件记录
        logger.info(f"[生成任务 {tender_generation_id}] 创建文件记录，文件ID: {file_id}")
        file_record = File(
            file_id=file_id,
            origin_name=f"采购征询文件_{uuid.uuid4().hex[:8]}.docx",
            file_name=file_name
        )
        db.add(file_record)
        db.flush()  # 刷新以获取 file_id
        
        # 更新生成采购征询文件记录
        tender_generation.file_id = file_id
        tender_generation.status = 'finished'
        db.commit()
        db.refresh(tender_generation)
        
        logger.info(f"[生成任务 {tender_generation_id}] 采购征询文件生成完成，文件ID: {file_id}, 生成记录ID: {tender_generation_id}")
        return file_record
        
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        logger.error(f"[生成任务 {tender_generation_id}] 生成采购征询文件失败: {str(e)}")
        logger.error(f"[生成任务 {tender_generation_id}] 错误堆栈:\n{error_trace}")
        
        db.rollback()
        # 更新状态为 failed
        try:
            tender_generation = db.query(TenderGeneration).filter(TenderGeneration.id == tender_generation_id).first()
            if tender_generation:
                tender_generation.status = 'failed'
                db.commit()
                logger.info(f"[生成任务 {tender_generation_id}] 已更新状态为 failed")
        except Exception as update_error:
            logger.error(f"[生成任务 {tender_generation_id}] 更新失败状态时出错: {str(update_error)}")
        raise
    finally:
        db.close()


@router.post("/generate-tender", summary="生成采购征询文件", description="根据业务需求文件和采购部门要求，使用LLM生成采购征询文件")
async def generate_tender(
    request: GenerateTenderRequest,
    db: Session = Depends(get_db)
):
    """
    生成采购征询文件接口
    
    此接口只触发生成动作，生成一个新的File就结束。生成完成后不会直接把征询文件ID给项目，
    需要使用 /tender-generation/{tender_generation_id}/use 接口来将文件关联到项目。
    一个项目可能有多次生成，可以通过 /tender-generation/{project_id}/list 接口查看所有生成记录。
    
    - **project_id**: 项目ID（必需）
    - **technical_requirement_file_id**: 业务需求文件ID
    - **procurement_requirement_file_id**: 采购部门要求文件ID
    
    返回：
    - **tender_generation_id**: 生成采购征询文件记录ID
    - **session_id**: 异步任务会话ID，可用于查询任务状态
    - **status**: 当前状态（running）
    - **message**: 提示信息
    """
    logger.info(f"[生成请求] 项目ID: {request.project_id}, 技术文件ID: {request.technical_requirement_file_id}, 采购文件ID: {request.procurement_requirement_file_id}")
    
    # 验证项目是否存在
    project = db.query(Project).filter(Project.id == request.project_id).first()
    if not project:
        logger.error(f"[生成请求] 项目不存在: {request.project_id}")
        raise HTTPException(
            status_code=404,
            detail=f"项目不存在: {request.project_id}"
        )
    
    logger.info(f"[生成请求] 项目验证成功: {project.name}")
    
    # 验证文件是否存在
    technical_file = db.query(File).filter(File.file_id == request.technical_requirement_file_id).first()
    if not technical_file:
        logger.error(f"[生成请求] 业务需求文件不存在: {request.technical_requirement_file_id}")
        raise HTTPException(
            status_code=404,
            detail=f"业务需求文件不存在: {request.technical_requirement_file_id}"
        )
    
    procurement_file = db.query(File).filter(File.file_id == request.procurement_requirement_file_id).first()
    if not procurement_file:
        logger.error(f"[生成请求] 采购部门要求文件不存在: {request.procurement_requirement_file_id}")
        raise HTTPException(
            status_code=404,
            detail=f"采购部门要求文件不存在: {request.procurement_requirement_file_id}"
        )
    
    logger.info(f"[生成请求] 文件验证成功 - 技术文件: {technical_file.origin_name}, 采购文件: {procurement_file.origin_name}")
    
    # 读取文件内容并构建 prompt
    files_folder = Path(settings.data_folder) / "files"
    technical_file_path = files_folder / technical_file.file_name
    procurement_file_path = files_folder / procurement_file.file_name
    
    if not technical_file_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"技术需求文件不存在: {technical_file_path}"
        )
    if not procurement_file_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"采购部门要求文件不存在: {procurement_file_path}"
        )
    
    # 读取文件内容
    logger.info(f"[生成请求] 开始读取文件内容")
    technical_file_ext = Path(technical_file.origin_name).suffix.lower()
    if technical_file_ext == '.docx':
        technical_content = read_docx_to_text(technical_file_path)
        logger.info(f"[生成请求] 技术文件读取成功，长度: {len(technical_content)} 字符")
    elif technical_file_ext == '.md':
        technical_content = read_markdown_file(technical_file_path)
        logger.info(f"[生成请求] 技术文件读取成功，长度: {len(technical_content)} 字符")
    else:
        logger.error(f"[生成请求] 不支持的技术文件格式: {technical_file_ext}")
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {technical_file_ext}"
        )
    
    procurement_file_ext = Path(procurement_file.origin_name).suffix.lower()
    if procurement_file_ext == '.docx':
        procurement_content = read_docx_to_text(procurement_file_path)
        logger.info(f"[生成请求] 采购文件读取成功，长度: {len(procurement_content)} 字符")
    elif procurement_file_ext == '.md':
        procurement_content = read_markdown_file(procurement_file_path)
        logger.info(f"[生成请求] 采购文件读取成功，长度: {len(procurement_content)} 字符")
    else:
        logger.error(f"[生成请求] 不支持的采购文件格式: {procurement_file_ext}")
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {procurement_file_ext}"
        )
    
    # 构建 prompt
    prompt = f"""你是一位专业的采购征询文件编写专家。请根据以下技术部门的需求文件和采购部门的要求，生成一份完整的、专业的采购征询文件。

## 技术部门需求文件

{technical_content}

## 采购部门要求

{procurement_content}

## 任务要求

请根据以上两个文件的内容，生成一份完整的采购征询文件（Markdown格式），因为后续会转换为docx格式，所以请不要包含任何格式标记。采购征询文件应包含以下章节结构：

1. **标题**：根据技术需求文件中的项目名称生成，格式为"[项目名称]采购征询文件"（居中显示）
2. **一、项目基本信息**：从技术需求文件中提取项目名称、项目编号等信息，从采购部门要求中提取采购单位、征询日期、投标截止时间等信息
3. **二、项目概述**：完整引用技术需求文件中的项目概述内容
4. **三、技术要求**：完整引用技术需求文件中的技术要求，包括所有子章节（如材料要求、制造工艺要求、质量要求、交付要求等）
5. **四、商务要求**：完整引用采购部门要求中的商务要求内容
6. **五、投标人资格要求**：完整引用采购部门要求中的投标人资格要求内容
7. **六、评分细则**：完整引用采购部门要求中的评分细则，并以表格形式呈现评分标准
8. **七、投标文件格式要求**：完整引用采购部门要求中的投标文件格式要求内容
9. **八、其他说明**：完整引用采购部门要求中的其他说明内容
10. **九、联系方式**：完整引用采购部门要求中的联系方式内容

## 注意事项

1. 保持专业、严谨的文档风格
2. 确保所有章节内容完整、逻辑清晰
3. 评分细则需要以表格形式呈现，表格应包含：评分项目、分值、评分标准、备注等列
4. 投标文件格式要求需要详细说明章节结构和格式评分标准
5. 使用标准的 Markdown 格式，包括标题（#）、列表（- 或 1.）、表格（|）等
6. 确保所有信息准确，不要遗漏任何重要内容
7. 项目名称、项目编号等具体信息应从技术需求文件中提取，不要自行编造
8. 采购单位、联系方式等信息应从采购部门要求中提取

请直接输出完整的采购征询文件 Markdown 内容，不要包含任何额外的说明或注释。"""
    
    # 使用 model_session 创建异步任务
    manager = get_manager()
    messages = [
        {
            "role": "system",
            "content": "你是一位专业的采购征询文件编写专家，擅长根据技术需求和采购要求编写完整、专业的采购征询文件。"
        },
        {
            "role": "user",
            "content": prompt
        }
    ]
    
    # 创建生成采购征询文件记录（状态为 running）
    logger.info(f"[生成请求] 创建生成记录")
    tender_generation = TenderGeneration(
        project_id=request.project_id,
        business_requirement_file_id=request.technical_requirement_file_id,
        procurement_requirement_file_id=request.procurement_requirement_file_id,
        status='running'
    )
    db.add(tender_generation)
    db.flush()  # 刷新以获取 ID
    tender_generation_id = tender_generation.id
    logger.info(f"[生成请求] 生成记录创建成功，ID: {tender_generation_id}")
    
    # 创建会话描述
    description = {
        "type": "generate_tender",
        "tender_generation_id": tender_generation_id,
        "project_id": request.project_id,
        "technical_file_id": request.technical_requirement_file_id,
        "procurement_file_id": request.procurement_requirement_file_id
    }
    
    # 创建并运行会话（异步任务）
    logger.info(f"[生成请求] 创建并启动LLM会话")
    manager = get_manager()
    session_id = await manager.create_and_run_session(
        messages=messages,
        model="deepseek-ai/DeepSeek-V3.2-Exp",
        description=description
    )
    logger.info(f"[生成请求] LLM会话创建成功，会话ID: {session_id}")
    
    # 更新生成采购征询文件记录的 model_session
    tender_generation.model_session = session_id
    db.commit()
    db.refresh(tender_generation)
    logger.info(f"[生成请求] 生成记录已更新会话ID")
    
    # 在后台任务中处理生成的文件保存（使用新线程执行异步任务）
    async def background_task():
        try:
            logger.info(f"[生成请求] 启动后台任务处理生成")
            await _generate_tender_task(
                tender_generation_id,
                request.technical_requirement_file_id,
                request.procurement_requirement_file_id,
                session_id
            )
        except Exception as e:
            logger.error(f"[生成请求] 后台任务执行失败: {str(e)}")
            import traceback
            logger.error(f"[生成请求] 后台任务错误堆栈:\n{traceback.format_exc()}")
    
    # 启动后台任务（不等待完成）
    asyncio.create_task(background_task())
    logger.info(f"[生成请求] 后台任务已启动，返回响应")
    
    return {
        "tender_generation_id": tender_generation_id,
        "session_id": session_id,
        "status": tender_generation.status,
        "message": "采购征询文件生成任务已启动，请使用 tender_generation_id 或 session_id 查询任务状态"
    }


@router.get("/tender-generation/{project_id}/list", summary="获取项目的采购征询文件生成列表", description="获取指定项目的所有采购征询文件生成记录列表")
async def get_tender_generation_list(
    project_id: int,
    db: Session = Depends(get_db)
):
    """
    获取项目的采购征询文件生成列表
    
    - **project_id**: 项目ID
    
    返回：
    - **list**: 生成记录列表，按时间倒序排列
      - **tender_generation_id**: 生成采购征询文件记录ID
      - **session_id**: 会话ID
      - **status**: 状态（running, finished, failed）
      - **file_id**: 生成的文件ID（如果已完成）
      - **file_name**: 生成的文件名称（如果已完成）
      - **created_at**: 创建时间（如果有）
    """
    # 验证项目是否存在
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(
            status_code=404,
            detail=f"项目不存在: {project_id}"
        )
    
    # 查询项目的所有生成记录，按ID倒序（最新的在前）
    tender_generations = db.query(TenderGeneration).filter(
        TenderGeneration.project_id == project_id
    ).order_by(TenderGeneration.id.desc()).all()
    
    result = []
    for tg in tender_generations:
        file_name = None
        if tg.file_id:
            file_record = db.query(File).filter(File.file_id == tg.file_id).first()
            if file_record:
                file_name = file_record.origin_name
        
        result.append({
            "tender_generation_id": tg.id,
            "session_id": tg.model_session,
            "status": tg.status,
            "file_id": tg.file_id,
            "file_name": file_name
        })
    
    return {
        "project_id": project_id,
        "list": result
    }


@router.get("/tender-generation/{tender_generation_id}/stream", summary="获取采购征询文件生成的流式响应", description="通过SSE流式返回采购征询文件生成内容")
async def get_tender_generation_stream(
    tender_generation_id: int,
    db: Session = Depends(get_db)
):
    """
    获取采购征询文件生成的流式响应
    
    根据表中对应的 model_session，从 model_manager 中获取 session，
    使用 get_response 进行流式返回。
    
    - **tender_generation_id**: 生成采购征询文件记录ID
    
    返回SSE流式数据
    """
    from fastapi.responses import StreamingResponse
    import json
    
    # 查询生成记录
    tender_generation = db.query(TenderGeneration).filter(
        TenderGeneration.id == tender_generation_id
    ).first()
    
    if not tender_generation:
        raise HTTPException(status_code=404, detail=f"生成采购征询文件记录不存在: {tender_generation_id}")
    
    # 如果没有session_id，无法获取流式响应
    if not tender_generation.model_session:
        raise HTTPException(status_code=400, detail="无法获取流式响应：会话ID不存在")
    
    session_id = tender_generation.model_session
    manager = get_manager()
    
    # 从 model_manager 中获取会话
    session = manager.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail=f"会话不存在: {session_id}")
    
    async def generate_stream():
        try:
            # 发送状态信息
            yield f"data: {json.dumps({'type': 'status', 'status': tender_generation.status})}\n\n"
            
            # 使用 get_response 进行流式返回
            async for chunk in session.get_response():
                if chunk.delta_content:
                    yield f"data: {json.dumps({'type': 'chunk', 'content': chunk.delta_content})}\n\n"
            
            # 发送完成状态
            # 重新查询以获取最新状态
            db.refresh(tender_generation)
            yield f"data: {json.dumps({'type': 'status', 'status': tender_generation.status})}\n\n"
            
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"
    
    return StreamingResponse(
        generate_stream(),
        media_type="text/event-stream"
    )


@router.post("/tender-generation/{tender_generation_id}/use", summary="使用生成的采购征询文件", description="将生成的采购征询文件关联到项目")
async def use_tender_generation(
    tender_generation_id: int,
    db: Session = Depends(get_db)
):
    """
    使用生成的采购征询文件
    
    将指定生成记录的文件ID关联到项目的 tender_document_file_id
    
    - **tender_generation_id**: 生成采购征询文件记录ID
    
    返回：
    - **message**: 提示信息
    - **project_id**: 项目ID
    - **file_id**: 关联的文件ID
    """
    # 查询生成记录
    tender_generation = db.query(TenderGeneration).filter(
        TenderGeneration.id == tender_generation_id
    ).first()
    
    if not tender_generation:
        raise HTTPException(
            status_code=404,
            detail=f"生成采购征询文件记录不存在: {tender_generation_id}"
        )
    
    # 检查状态是否为已完成
    if tender_generation.status != 'finished':
        raise HTTPException(
            status_code=400,
            detail=f"生成记录状态为 {tender_generation.status}，只有已完成（finished）的记录才能使用"
        )
    
    # 检查是否有文件ID
    if not tender_generation.file_id:
        raise HTTPException(
            status_code=400,
            detail="生成记录没有关联的文件ID"
        )
    
    # 验证文件是否存在
    file_record = db.query(File).filter(File.file_id == tender_generation.file_id).first()
    if not file_record:
        raise HTTPException(
            status_code=404,
            detail=f"文件不存在: {tender_generation.file_id}"
        )
    
    # 获取项目
    project = db.query(Project).filter(Project.id == tender_generation.project_id).first()
    if not project:
        raise HTTPException(
            status_code=404,
            detail=f"项目不存在: {tender_generation.project_id}"
        )
    
    # 将文件ID关联到项目
    project.tender_document_file_id = tender_generation.file_id
    db.commit()
    db.refresh(project)
    
    return {
        "message": "采购征询文件已成功关联到项目",
        "project_id": project.id,
        "file_id": tender_generation.file_id,
        "file_name": file_record.origin_name
    }


@router.delete("/tender-generation/{tender_generation_id}", summary="删除采购征询文件生成记录", description="删除指定的采购征询文件生成记录及其关联的文件")
async def delete_tender_generation(
    tender_generation_id: int,
    db: Session = Depends(get_db)
):
    """
    删除采购征询文件生成记录
    
    删除记录时，如果有关联的文件，也会删除文件记录和物理文件。
    
    - **tender_generation_id**: 生成采购征询文件记录ID
    
    返回：
    - **message**: 提示信息
    - **deleted_file**: 是否删除了关联的文件
    """
    # 查询生成记录
    tender_generation = db.query(TenderGeneration).filter(
        TenderGeneration.id == tender_generation_id
    ).first()
    
    if not tender_generation:
        raise HTTPException(
            status_code=404,
            detail=f"生成采购征询文件记录不存在: {tender_generation_id}"
        )
    
    deleted_file = False
    file_id_to_delete = None
    
    # 如果有关联的文件，先获取文件ID
    if tender_generation.file_id:
        file_id_to_delete = tender_generation.file_id
        # 查询文件记录
        file_record = db.query(File).filter(File.file_id == file_id_to_delete).first()
        
        if file_record:
            # 删除物理文件
            try:
                files_folder = Path(settings.data_folder) / "files"
                file_path = files_folder / file_record.file_name
                if file_path.exists():
                    file_path.unlink()
                    deleted_file = True
            except Exception as e:
                print(f"删除物理文件失败: {str(e)}")
                # 继续删除记录，不因为文件删除失败而中断
            
            # 删除文件记录
            try:
                db.delete(file_record)
                db.flush()  # 刷新以确保文件记录被删除
            except Exception as e:
                print(f"删除文件记录失败: {str(e)}")
                db.rollback()
                raise HTTPException(
                    status_code=500,
                    detail=f"删除文件记录失败: {str(e)}"
                )
    
    # 删除生成记录
    try:
        db.delete(tender_generation)
        db.commit()
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail=f"删除生成记录失败: {str(e)}"
        )
    
    return {
        "message": "采购征询文件生成记录已成功删除",
        "tender_generation_id": tender_generation_id,
        "deleted_file": deleted_file
    }

