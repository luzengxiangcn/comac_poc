"""
生成飞机机身制造相关的招标文件和供应商标书的 Mock 数据
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import json


def create_technical_requirement_document():
    """创建技术部门需求文件（docx格式）
    
    该文件包含技术部门提出的项目需求，包括：
    - 项目基本信息
    - 项目概述
    - 技术要求（材料、工艺、质量、交付等）
    """
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('C919大型客机机身段制造项目技术需求文件', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 文档说明
    doc.add_paragraph('本文档由技术部门编制，用于明确项目技术需求，作为招标文件编制的技术依据。')
    
    # 项目基本信息
    doc.add_heading('一、项目基本信息', 1)
    p = doc.add_paragraph()
    p.add_run('项目名称：').bold = True
    p.add_run('C919大型客机机身段制造项目')
    
    p = doc.add_paragraph()
    p.add_run('项目编号：').bold = True
    p.add_run('COMAC-2024-FUSELAGE-001')
    
    p = doc.add_paragraph()
    p.add_run('需求部门：').bold = True
    p.add_run('技术部门')
    
    p = doc.add_paragraph()
    p.add_run('编制日期：').bold = True
    p.add_run('2024年3月10日')
    
    # 项目概述
    doc.add_heading('二、项目概述', 1)
    doc.add_paragraph(
        '本项目旨在采购C919大型客机机身中段和后段的制造服务，包括但不限于：'
    )
    doc.add_paragraph('1. 机身中段（含客舱段）制造，长度约15米，直径约3.96米', style='List Bullet')
    doc.add_paragraph('2. 机身后段（含尾锥段）制造，长度约8米', style='List Bullet')
    doc.add_paragraph('3. 机身蒙皮、框架、桁条等结构件的加工与装配', style='List Bullet')
    doc.add_paragraph('4. 机身段总装、检测与交付', style='List Bullet')
    doc.add_paragraph('5. 相关技术文档和工艺文件的提供', style='List Bullet')
    
    # 技术要求
    doc.add_heading('三、技术要求', 1)
    doc.add_heading('3.1 材料要求', 2)
    doc.add_paragraph('1. 蒙皮材料：采用航空级铝合金2024-T3或7075-T6，厚度范围1.2-3.0mm', style='List Number')
    doc.add_paragraph('2. 框架材料：采用航空级铝合金或钛合金，满足AS9100标准', style='List Number')
    doc.add_paragraph('3. 紧固件：符合NAS标准，提供材料证明和检测报告', style='List Number')
    doc.add_paragraph('4. 所有材料需提供完整的追溯性文件，包括批次号、供应商信息等', style='List Number')
    
    doc.add_heading('3.2 制造工艺要求', 2)
    doc.add_paragraph('1. 蒙皮成形：采用数控拉伸成形或滚弯成形，表面粗糙度Ra≤1.6μm', style='List Number')
    doc.add_paragraph('2. 机械加工：采用五轴数控加工中心，加工精度±0.05mm', style='List Number')
    doc.add_paragraph('3. 装配工艺：采用数字化装配技术，装配精度满足图纸要求', style='List Number')
    doc.add_paragraph('4. 焊接工艺：如涉及焊接，需符合AWS D17.1标准，焊工需持证上岗', style='List Number')
    doc.add_paragraph('5. 表面处理：需进行阳极氧化或喷涂处理，满足防腐要求', style='List Number')
    
    doc.add_heading('3.3 质量要求', 2)
    doc.add_paragraph('1. 质量管理体系：供应商需通过AS9100质量管理体系认证', style='List Number')
    doc.add_paragraph('2. 检测要求：所有关键尺寸需100%检测，并提供检测报告', style='List Number')
    doc.add_paragraph('3. 无损检测：关键部位需进行X射线、超声波或渗透检测', style='List Number')
    doc.add_paragraph('4. 交付标准：产品需满足CCAR-25部适航要求', style='List Number')
    doc.add_paragraph('5. 不合格品处理：建立不合格品控制程序，确保不合格品不流入下道工序', style='List Number')
    
    doc.add_heading('3.4 交付要求', 2)
    doc.add_paragraph('1. 交付时间：合同签订后18个月内完成首批交付', style='List Number')
    doc.add_paragraph('2. 交付地点：招标方指定地点（上海浦东）', style='List Number')
    doc.add_paragraph('3. 交付内容：包括产品、技术文档、检测报告、质量证明文件等', style='List Number')
    doc.add_paragraph('4. 包装运输：采用专用工装和包装，确保运输过程中产品不受损', style='List Number')
    
    # 保存文件
    output_path = Path(__file__).parent / '需求文件' / '技术需求文件_C919机身段制造项目.docx'
    output_path.parent.mkdir(exist_ok=True)
    doc.save(str(output_path))
    print(f'技术需求文件已生成：{output_path}')
    return output_path


def create_procurement_requirement_document():
    """创建采购部门标书要求文件（docx格式）
    
    该文件包含采购部门对标书格式、内容、评分标准等方面的统一要求，包括：
    - 商务要求
    - 投标人资格要求
    - 评分细则
    - 标书格式要求
    - 其他说明
    - 联系方式
    """
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('采购部门标书要求规范', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 文档说明
    doc.add_paragraph('本文档规定了采购部门对标书格式、内容、评分标准等方面的统一要求。所有招标文件必须包含以下内容。')
    
    # 商务要求
    doc.add_heading('一、商务要求', 1)
    doc.add_paragraph('招标文件必须包含以下商务条款：')
    doc.add_paragraph('1. 投标保证金：人民币100万元，投标截止前到账', style='List Number')
    doc.add_paragraph('2. 付款方式：合同签订后支付30%预付款，交付验收合格后支付65%，质保期满后支付5%', style='List Number')
    doc.add_paragraph('3. 质保期：自交付验收合格之日起36个月', style='List Number')
    doc.add_paragraph('4. 违约责任：延期交付按合同金额每日0.5‰支付违约金', style='List Number')
    doc.add_paragraph('5. 知识产权：投标方需保证不侵犯第三方知识产权，并承担相应责任', style='List Number')
    
    # 投标人资格要求
    doc.add_heading('二、投标人资格要求', 1)
    doc.add_paragraph('招标文件必须明确以下投标人资格要求：')
    doc.add_paragraph('1. 具有独立法人资格，注册资金不低于5000万元人民币', style='List Number')
    doc.add_paragraph('2. 具有航空制造相关资质，近三年内承担过类似项目', style='List Number')
    doc.add_paragraph('3. 具有完善的质量管理体系，通过AS9100认证', style='List Number')
    doc.add_paragraph('4. 财务状况良好，无重大诉讼和不良信用记录', style='List Number')
    doc.add_paragraph('5. 具有相应的技术团队和生产设备', style='List Number')
    doc.add_paragraph('6. 未被列入失信被执行人名单', style='List Number')
    
    # 评分细则
    doc.add_heading('三、评分细则', 1)
    doc.add_paragraph('招标文件必须包含详细的评分细则，采用综合评分法，总分100分。评分标准如下：')
    
    doc.add_heading('3.1 技术方案（30分）', 2)
    p = doc.add_paragraph()
    p.add_run('制造工艺方案（12分）：').bold = True
    p.add_run('工艺路线清晰、技术先进、可操作性强')
    
    p = doc.add_paragraph()
    p.add_run('质量控制方案（10分）：').bold = True
    p.add_run('质量体系完善、检测手段齐全、追溯性强')
    
    p = doc.add_paragraph()
    p.add_run('技术团队能力（8分）：').bold = True
    p.add_run('团队经验丰富、技术实力强')
    
    doc.add_heading('3.2 商务报价（40分）', 2)
    p = doc.add_paragraph()
    p.add_run('总报价（30分）：').bold = True
    p.add_run('总价最低者得30分，其他按公式计算（公式：(最低价/投标价)×30）')
    
    p = doc.add_paragraph()
    p.add_run('付款方式（10分）：').bold = True
    p.add_run('付款条件是否有利于招标方')
    
    doc.add_heading('3.3 企业资质（15分）', 2)
    p = doc.add_paragraph()
    p.add_run('资质认证（5分）：').bold = True
    p.add_run('AS9100认证、适航认证等')
    
    p = doc.add_paragraph()
    p.add_run('类似业绩（10分）：').bold = True
    p.add_run('近三年类似项目经验，每个项目2分，最高10分')
    
    doc.add_heading('3.4 交付能力（10分）', 2)
    p = doc.add_paragraph()
    p.add_run('生产设备（5分）：').bold = True
    p.add_run('设备先进、满足生产需求')
    
    p = doc.add_paragraph()
    p.add_run('交付计划（5分）：').bold = True
    p.add_run('计划合理、可执行性强')
    
    doc.add_heading('3.5 服务保障（5分）', 2)
    doc.add_paragraph('售后服务、技术支持、培训')
    
    # 标书格式要求
    doc.add_heading('四、标书格式要求', 1)
    doc.add_paragraph('招标文件必须明确标书格式要求，包括：')
    
    doc.add_heading('4.1 章节结构要求', 2)
    doc.add_paragraph('投标文件必须包含以下章节，且章节顺序不得更改：')
    
    doc.add_paragraph('1. 一、投标函', style='List Number')
    doc.add_paragraph('   必须包含：投标总价、交付周期、质保期、有效期、法定代表人签字', style='List Bullet')
    
    doc.add_paragraph('2. 二、企业概况', style='List Number')
    doc.add_paragraph('   2.1 企业简介', style='List Bullet')
    doc.add_paragraph('   2.2 资质认证', style='List Bullet')
    doc.add_paragraph('   必须包含：企业名称、注册地址、注册资本、成立时间、社会信用代码、法定代表人', style='List Bullet')
    
    doc.add_paragraph('3. 三、技术方案', style='List Number')
    doc.add_paragraph('   3.1 制造工艺方案', style='List Bullet')
    doc.add_paragraph('   3.2 质量控制方案', style='List Bullet')
    doc.add_paragraph('   3.3 技术团队', style='List Bullet')
    
    doc.add_paragraph('4. 四、类似业绩', style='List Number')
    doc.add_paragraph('   必须包含：项目名称、项目时间、项目内容、合同金额', style='List Bullet')
    
    doc.add_paragraph('5. 五、生产设备', style='List Number')
    doc.add_paragraph('   必须包含：设备名称、型号规格、数量、制造商', style='List Bullet')
    
    doc.add_paragraph('6. 六、交付计划', style='List Number')
    doc.add_paragraph('   必须包含：各阶段工作内容和完成时间', style='List Bullet')
    
    doc.add_paragraph('7. 七、商务报价', style='List Number')
    doc.add_paragraph('   必须包含：报价明细、总报价、付款方式', style='List Bullet')
    
    doc.add_paragraph('8. 八、服务保障', style='List Number')
    doc.add_paragraph('   必须包含：售后服务、技术支持、培训服务', style='List Bullet')
    
    doc.add_paragraph('9. 九、附件清单', style='List Number')
    
    p = doc.add_paragraph()
    p.add_run('注：').bold = True
    p.add_run('缺少任何章节或章节顺序不符合要求的，视为格式不符合要求。')
    
    doc.add_heading('4.2 格式评分', 2)
    doc.add_paragraph('标书格式符合性将作为评分项之一，总分5分：')
    doc.add_paragraph('• 完全符合格式要求（所有章节齐全、顺序正确、内容完整）：5分', style='List Bullet')
    doc.add_paragraph('• 基本符合，有轻微不符合项：3-4分', style='List Bullet')
    doc.add_paragraph('• 部分符合，有明显不符合项（如缺少重要章节）：1-2分', style='List Bullet')
    doc.add_paragraph('• 严重不符合格式要求：0分，且可能被判定为无效投标', style='List Bullet')
    
    # 其他说明
    doc.add_heading('五、其他说明', 1)
    doc.add_paragraph('招标文件应包含以下其他说明：')
    doc.add_paragraph('1. 投标文件需密封提交，一式五份（正本一份，副本四份）', style='List Number')
    doc.add_paragraph('2. 投标文件应包括：投标函、技术方案、商务报价、资质证明、业绩证明等', style='List Number')
    doc.add_paragraph('3. 开标时间：2024年5月5日 9:00（北京时间）', style='List Number')
    doc.add_paragraph('4. 评标结果将在开标后15个工作日内公布', style='List Number')
    doc.add_paragraph('5. 招标方保留对招标文件的解释权', style='List Number')
    
    # 联系方式
    doc.add_heading('六、联系方式', 1)
    doc.add_paragraph('招标文件必须包含以下联系方式：')
    p = doc.add_paragraph()
    p.add_run('招标单位：').bold = True
    p.add_run('中国商用飞机有限责任公司')
    
    p = doc.add_paragraph()
    p.add_run('联系人：').bold = True
    p.add_run('张工程师')
    
    p = doc.add_paragraph()
    p.add_run('联系电话：').bold = True
    p.add_run('021-12345678')
    
    p = doc.add_paragraph()
    p.add_run('电子邮箱：').bold = True
    p.add_run('tender@comac.cc')
    
    p = doc.add_paragraph()
    p.add_run('地址：').bold = True
    p.add_run('上海市浦东新区上飞路919号')
    
    # 招标文件基本信息
    doc.add_heading('七、招标文件基本信息', 1)
    doc.add_paragraph('招标文件应包含以下基本信息（由采购部门填写）：')
    p = doc.add_paragraph()
    p.add_run('招标单位：').bold = True
    p.add_run('中国商用飞机有限责任公司')
    
    p = doc.add_paragraph()
    p.add_run('招标日期：').bold = True
    p.add_run('2024年3月15日')
    
    p = doc.add_paragraph()
    p.add_run('投标截止时间：').bold = True
    p.add_run('2024年4月30日 17:00（北京时间）')
    
    # 说明
    doc.add_heading('说明', 1)
    doc.add_paragraph('本文档为采购部门内部规范，用于指导招标文件的编制。招标文件应结合技术部门的技术需求文件和本文档的要求进行编制。')
    
    # 保存文件
    output_path = Path(__file__).parent / '需求文件' / '采购部门标书要求.docx'
    output_path.parent.mkdir(exist_ok=True)
    doc.save(str(output_path))
    print(f'采购部门标书要求文件已生成：{output_path}')
    return output_path


def create_tender_document():
    """创建飞机机身制造招标文件
    
    本函数根据以下两个来源组合生成招标文件：
    1. 技术部门需求文件（需求文件/技术需求文件_C919机身段制造项目.docx）
       - 包含：项目基本信息、项目概述、技术要求
    2. 采购部门标书要求（需求文件/采购部门标书要求.docx）
       - 包含：商务要求、投标人资格要求、评分细则、标书格式要求、其他说明、联系方式
    
    招标文件的结构：
    - 一、项目基本信息（来自技术需求 + 采购部门补充的招标信息）
    - 二、项目概述（来自技术需求）
    - 三、技术要求（来自技术需求）
    - 四、商务要求（来自采购部门要求）
    - 五、投标人资格要求（来自采购部门要求）
    - 六、评分细则（来自采购部门要求）
    - 七、标书格式要求（来自采购部门要求）
    - 八、其他说明（来自采购部门要求）
    - 九、联系方式（来自采购部门要求）
    """
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 标题
    title = doc.add_heading('飞机机身制造项目招标文件', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 项目信息
    doc.add_heading('一、项目基本信息', 1)
    p = doc.add_paragraph()
    p.add_run('项目名称：').bold = True
    p.add_run('C919大型客机机身段制造项目')
    
    p = doc.add_paragraph()
    p.add_run('项目编号：').bold = True
    p.add_run('COMAC-2024-FUSELAGE-001')
    
    p = doc.add_paragraph()
    p.add_run('招标单位：').bold = True
    p.add_run('中国商用飞机有限责任公司')
    
    p = doc.add_paragraph()
    p.add_run('招标日期：').bold = True
    p.add_run('2024年3月15日')
    
    p = doc.add_paragraph()
    p.add_run('投标截止时间：').bold = True
    p.add_run('2024年4月30日 17:00（北京时间）')
    
    # 项目概述
    doc.add_heading('二、项目概述', 1)
    doc.add_paragraph(
        '本项目旨在采购C919大型客机机身中段和后段的制造服务，包括但不限于：'
    )
    doc.add_paragraph('1. 机身中段（含客舱段）制造，长度约15米，直径约3.96米', style='List Bullet')
    doc.add_paragraph('2. 机身后段（含尾锥段）制造，长度约8米', style='List Bullet')
    doc.add_paragraph('3. 机身蒙皮、框架、桁条等结构件的加工与装配', style='List Bullet')
    doc.add_paragraph('4. 机身段总装、检测与交付', style='List Bullet')
    doc.add_paragraph('5. 相关技术文档和工艺文件的提供', style='List Bullet')
    
    # 技术要求
    doc.add_heading('三、技术要求', 1)
    doc.add_heading('3.1 材料要求', 2)
    doc.add_paragraph('1. 蒙皮材料：采用航空级铝合金2024-T3或7075-T6，厚度范围1.2-3.0mm', style='List Number')
    doc.add_paragraph('2. 框架材料：采用航空级铝合金或钛合金，满足AS9100标准', style='List Number')
    doc.add_paragraph('3. 紧固件：符合NAS标准，提供材料证明和检测报告', style='List Number')
    doc.add_paragraph('4. 所有材料需提供完整的追溯性文件，包括批次号、供应商信息等', style='List Number')
    
    doc.add_heading('3.2 制造工艺要求', 2)
    doc.add_paragraph('1. 蒙皮成形：采用数控拉伸成形或滚弯成形，表面粗糙度Ra≤1.6μm', style='List Number')
    doc.add_paragraph('2. 机械加工：采用五轴数控加工中心，加工精度±0.05mm', style='List Number')
    doc.add_paragraph('3. 装配工艺：采用数字化装配技术，装配精度满足图纸要求', style='List Number')
    doc.add_paragraph('4. 焊接工艺：如涉及焊接，需符合AWS D17.1标准，焊工需持证上岗', style='List Number')
    doc.add_paragraph('5. 表面处理：需进行阳极氧化或喷涂处理，满足防腐要求', style='List Number')
    
    doc.add_heading('3.3 质量要求', 2)
    doc.add_paragraph('1. 质量管理体系：供应商需通过AS9100质量管理体系认证', style='List Number')
    doc.add_paragraph('2. 检测要求：所有关键尺寸需100%检测，并提供检测报告', style='List Number')
    doc.add_paragraph('3. 无损检测：关键部位需进行X射线、超声波或渗透检测', style='List Number')
    doc.add_paragraph('4. 交付标准：产品需满足CCAR-25部适航要求', style='List Number')
    doc.add_paragraph('5. 不合格品处理：建立不合格品控制程序，确保不合格品不流入下道工序', style='List Number')
    
    doc.add_heading('3.4 交付要求', 2)
    doc.add_paragraph('1. 交付时间：合同签订后18个月内完成首批交付', style='List Number')
    doc.add_paragraph('2. 交付地点：招标方指定地点（上海浦东）', style='List Number')
    doc.add_paragraph('3. 交付内容：包括产品、技术文档、检测报告、质量证明文件等', style='List Number')
    doc.add_paragraph('4. 包装运输：采用专用工装和包装，确保运输过程中产品不受损', style='List Number')
    
    # 商务要求
    doc.add_heading('四、商务要求', 1)
    doc.add_paragraph('1. 投标保证金：人民币100万元，投标截止前到账', style='List Number')
    doc.add_paragraph('2. 付款方式：合同签订后支付30%预付款，交付验收合格后支付65%，质保期满后支付5%', style='List Number')
    doc.add_paragraph('3. 质保期：自交付验收合格之日起36个月', style='List Number')
    doc.add_paragraph('4. 违约责任：延期交付按合同金额每日0.5‰支付违约金', style='List Number')
    doc.add_paragraph('5. 知识产权：投标方需保证不侵犯第三方知识产权，并承担相应责任', style='List Number')
    
    # 投标人资格要求
    doc.add_heading('五、投标人资格要求', 1)
    doc.add_paragraph('1. 具有独立法人资格，注册资金不低于5000万元人民币', style='List Number')
    doc.add_paragraph('2. 具有航空制造相关资质，近三年内承担过类似项目', style='List Number')
    doc.add_paragraph('3. 具有完善的质量管理体系，通过AS9100认证', style='List Number')
    doc.add_paragraph('4. 财务状况良好，无重大诉讼和不良信用记录', style='List Number')
    doc.add_paragraph('5. 具有相应的技术团队和生产设备', style='List Number')
    doc.add_paragraph('6. 未被列入失信被执行人名单', style='List Number')
    
    # 评分细则
    doc.add_heading('六、评分细则', 1)
    doc.add_paragraph('本次评标采用综合评分法，总分100分，具体评分标准如下：')
    
    # 评分表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # 表头
    header_cells = table.rows[0].cells
    header_cells[0].text = '评分项目'
    header_cells[1].text = '分值'
    header_cells[2].text = '评分标准'
    header_cells[3].text = '备注'
    
    # 技术方案评分
    row = table.add_row()
    row.cells[0].text = '技术方案（30分）'
    row.cells[1].text = '30'
    row.cells[2].text = '工艺方案合理性、技术先进性、可行性'
    row.cells[3].text = '由技术专家评分'
    
    # 技术方案子项
    row = table.add_row()
    row.cells[0].text = '   - 制造工艺方案'
    row.cells[1].text = '12'
    row.cells[2].text = '工艺路线清晰、技术先进、可操作性强'
    row.cells[3].text = ''
    
    row = table.add_row()
    row.cells[0].text = '   - 质量控制方案'
    row.cells[1].text = '10'
    row.cells[2].text = '质量体系完善、检测手段齐全、追溯性强'
    row.cells[3].text = ''
    
    row = table.add_row()
    row.cells[0].text = '   - 技术团队能力'
    row.cells[1].text = '8'
    row.cells[2].text = '团队经验丰富、技术实力强'
    row.cells[3].text = ''
    
    # 商务报价评分
    row = table.add_row()
    row.cells[0].text = '商务报价（40分）'
    row.cells[1].text = '40'
    row.cells[2].text = '价格合理性、性价比'
    row.cells[3].text = '最低价得满分，其他按比例'
    
    row = table.add_row()
    row.cells[0].text = '   - 总报价'
    row.cells[1].text = '30'
    row.cells[2].text = '总价最低者得30分，其他按公式计算'
    row.cells[3].text = '公式：(最低价/投标价)×30'
    
    row = table.add_row()
    row.cells[0].text = '   - 付款方式'
    row.cells[1].text = '10'
    row.cells[2].text = '付款条件是否有利于招标方'
    row.cells[3].text = ''
    
    # 企业资质评分
    row = table.add_row()
    row.cells[0].text = '企业资质（15分）'
    row.cells[1].text = '15'
    row.cells[2].text = '企业实力、资质认证、业绩'
    row.cells[3].text = ''
    
    row = table.add_row()
    row.cells[0].text = '   - 资质认证'
    row.cells[1].text = '5'
    row.cells[2].text = 'AS9100认证、适航认证等'
    row.cells[3].text = ''
    
    row = table.add_row()
    row.cells[0].text = '   - 类似业绩'
    row.cells[1].text = '10'
    row.cells[2].text = '近三年类似项目经验，每个项目2分，最高10分'
    row.cells[3].text = ''
    
    # 交付能力评分
    row = table.add_row()
    row.cells[0].text = '交付能力（10分）'
    row.cells[1].text = '10'
    row.cells[2].text = '生产设备、产能、交付计划'
    row.cells[3].text = ''
    
    row = table.add_row()
    row.cells[0].text = '   - 生产设备'
    row.cells[1].text = '5'
    row.cells[2].text = '设备先进、满足生产需求'
    row.cells[3].text = ''
    
    row = table.add_row()
    row.cells[0].text = '   - 交付计划'
    row.cells[1].text = '5'
    row.cells[2].text = '计划合理、可执行性强'
    row.cells[3].text = ''
    
    # 服务保障评分
    row = table.add_row()
    row.cells[0].text = '服务保障（5分）'
    row.cells[1].text = '5'
    row.cells[2].text = '售后服务、技术支持、培训'
    row.cells[3].text = ''
    
    # 标书格式要求
    doc.add_heading('七、标书格式要求', 1)
    doc.add_paragraph('投标文件必须严格按照以下格式要求编制，不符合格式要求的投标文件将被视为无效投标。')
    
    doc.add_heading('7.1 章节结构要求', 2)
    doc.add_paragraph('投标文件必须包含以下章节，且章节顺序不得更改：')
    doc.add_paragraph('一、投标函', style='List Bullet')
    doc.add_paragraph('   必须包含：投标总价、交付周期、质保期、有效期、法定代表人签字', style='List Bullet')
    doc.add_paragraph('二、企业概况', style='List Bullet')
    doc.add_paragraph('   2.1 企业简介', style='List Bullet')
    doc.add_paragraph('   2.2 资质认证', style='List Bullet')
    doc.add_paragraph('   必须包含：企业名称、注册地址、注册资本、成立时间、社会信用代码、法定代表人', style='List Bullet')
    doc.add_paragraph('三、技术方案', style='List Bullet')
    doc.add_paragraph('   3.1 制造工艺方案', style='List Bullet')
    doc.add_paragraph('   3.2 质量控制方案', style='List Bullet')
    doc.add_paragraph('   3.3 技术团队', style='List Bullet')
    doc.add_paragraph('四、类似业绩', style='List Bullet')
    doc.add_paragraph('   必须包含：项目名称、项目时间、项目内容、合同金额', style='List Bullet')
    doc.add_paragraph('五、生产设备', style='List Bullet')
    doc.add_paragraph('   必须包含：设备名称、型号规格、数量、制造商', style='List Bullet')
    doc.add_paragraph('六、交付计划', style='List Bullet')
    doc.add_paragraph('   必须包含：各阶段工作内容和完成时间', style='List Bullet')
    doc.add_paragraph('七、商务报价', style='List Bullet')
    doc.add_paragraph('   必须包含：报价明细、总报价、付款方式', style='List Bullet')
    doc.add_paragraph('八、服务保障', style='List Bullet')
    doc.add_paragraph('   必须包含：售后服务、技术支持、培训服务', style='List Bullet')
    doc.add_paragraph('九、附件清单', style='List Bullet')
    doc.add_paragraph('注：缺少任何章节或章节顺序不符合要求的，视为格式不符合要求。', style='List Number')
    
    doc.add_heading('7.2 格式评分', 2)
    doc.add_paragraph('标书格式符合性将作为评分项之一，总分5分：')
    doc.add_paragraph('• 完全符合格式要求（所有章节齐全、顺序正确、内容完整）：5分', style='List Bullet')
    doc.add_paragraph('• 基本符合，有轻微不符合项：3-4分', style='List Bullet')
    doc.add_paragraph('• 部分符合，有明显不符合项（如缺少重要章节）：1-2分', style='List Bullet')
    doc.add_paragraph('• 严重不符合格式要求：0分，且可能被判定为无效投标', style='List Bullet')
    
    # 其他说明
    doc.add_heading('八、其他说明', 1)
    doc.add_paragraph('1. 投标文件需密封提交，一式五份（正本一份，副本四份）', style='List Number')
    doc.add_paragraph('2. 投标文件应包括：投标函、技术方案、商务报价、资质证明、业绩证明等', style='List Number')
    doc.add_paragraph('3. 开标时间：2024年5月5日 9:00（北京时间）', style='List Number')
    doc.add_paragraph('4. 评标结果将在开标后15个工作日内公布', style='List Number')
    doc.add_paragraph('5. 招标方保留对招标文件的解释权', style='List Number')
    
    # 联系方式
    doc.add_heading('九、联系方式', 1)
    doc.add_paragraph('招标单位：中国商用飞机有限责任公司')
    doc.add_paragraph('联系人：张工程师')
    doc.add_paragraph('联系电话：021-12345678')
    doc.add_paragraph('电子邮箱：tender@comac.cc')
    doc.add_paragraph('地址：上海市浦东新区上飞路919号')
    
    # 保存文件
    output_path = Path(__file__).parent / '飞机机身制造项目招标文件.docx'
    doc.save(str(output_path))
    print(f'招标文件已生成：{output_path}')
    return output_path


def create_supplier_bid_document(supplier_name, supplier_info, format_compliant=True):
    """创建供应商投标文件
    
    Args:
        supplier_name: 供应商名称
        supplier_info: 供应商信息字典
        format_compliant: 是否符合格式要求，True为符合，False为不符合
    """
    doc = Document()
    
    # 设置页面格式
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4高度
    section.page_width = Inches(8.27)    # A4宽度
    section.top_margin = Inches(0.98)    # 2.5cm
    section.bottom_margin = Inches(0.98)
    section.left_margin = Inches(1.18)   # 3cm
    section.right_margin = Inches(1.18)
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # 设置标题样式
    heading1_style = doc.styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = '黑体'
    heading1_font.size = Pt(16)
    heading1_font.bold = True
    
    heading2_style = doc.styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = '黑体'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    
    # 标题
    title = doc.add_heading(f'{supplier_name}投标文件', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('C919大型客机机身段制造项目')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0] if subtitle.runs else subtitle.add_run()
    subtitle_format.font.size = Pt(14)
    subtitle_format.bold = True
    
    
    # 投标函
    doc.add_heading('一、投标函', 1)
    doc.add_paragraph(
        f'致：中国商用飞机有限责任公司\n\n'
        f'我公司（{supplier_name}）经认真研究招标文件，决定参加"C919大型客机机身段制造项目"的投标。\n\n'
        f'我公司承诺：\n'
        f'1. 完全理解并接受招标文件的所有条款和要求\n'
        f'2. 投标总价为人民币 {supplier_info["total_price"]} 万元\n'
        f'3. 交付周期：{supplier_info["delivery_period"]} 个月\n'
        f'4. 质保期：{supplier_info["warranty_period"]} 个月\n'
        f'5. 本投标文件自开标之日起90天内有效\n\n'
        f'投标人：{supplier_name}\n'
        f'法定代表人（签字）：\n'
        f'日期：2024年4月25日'
    )
    
    # 企业概况
    doc.add_heading('二、企业概况', 1)
    doc.add_paragraph(f'企业名称：{supplier_name}')
    doc.add_paragraph(f'注册地址：{supplier_info["address"]}')
    doc.add_paragraph(f'注册资本：{supplier_info["registered_capital"]} 万元')
    doc.add_paragraph(f'成立时间：{supplier_info["established_date"]}')
    doc.add_paragraph(f'社会信用代码：{supplier_info["registration_number"]}')
    doc.add_paragraph(f'法定代表人：{supplier_info["legal_representative"]}')
    
    doc.add_heading('2.1 企业简介', 2)
    doc.add_paragraph(supplier_info["company_intro"])
    
    doc.add_heading('2.2 资质认证', 2)
    for cert in supplier_info["certifications"]:
        doc.add_paragraph(f'• {cert}', style='List Bullet')
    
    # 技术方案
    doc.add_heading('三、技术方案', 1)
    doc.add_heading('3.1 制造工艺方案', 2)
    doc.add_paragraph(supplier_info["manufacturing_process"])
    
    doc.add_heading('3.2 质量控制方案', 2)
    doc.add_paragraph('1. 质量管理体系：')
    doc.add_paragraph(supplier_info["quality_system"], style='List Bullet')
    
    doc.add_paragraph('2. 检测设备：')
    for equipment in supplier_info["testing_equipment"]:
        doc.add_paragraph(f'• {equipment}', style='List Bullet')
    
    doc.add_paragraph('3. 关键控制点：')
    for control_point in supplier_info["control_points"]:
        doc.add_paragraph(f'• {control_point}', style='List Bullet')
    
    doc.add_heading('3.3 技术团队', 2)
    doc.add_paragraph(f'项目技术团队共 {supplier_info["team_size"]} 人，其中：')
    doc.add_paragraph(f'• 高级工程师：{supplier_info["senior_engineers"]} 人', style='List Bullet')
    doc.add_paragraph(f'• 工程师：{supplier_info["engineers"]} 人', style='List Bullet')
    doc.add_paragraph(f'• 技术员：{supplier_info["technicians"]} 人', style='List Bullet')
    doc.add_paragraph(f'• 项目负责人：{supplier_info["project_manager"]}', style='List Bullet')
    
    # 类似业绩
    doc.add_heading('四、类似业绩', 1)
    for i, project in enumerate(supplier_info["similar_projects"], 1):
        doc.add_paragraph(f'{i}. {project["name"]}', style='List Number')
        doc.add_paragraph(f'   项目时间：{project["period"]}')
        doc.add_paragraph(f'   项目内容：{project["content"]}')
        doc.add_paragraph(f'   合同金额：{project["amount"]} 万元')
    
    # 生产设备
    doc.add_heading('五、生产设备', 1)
    doc.add_paragraph('主要生产设备清单：')
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = '设备名称'
    header_cells[1].text = '型号规格'
    header_cells[2].text = '数量'
    header_cells[3].text = '制造商'
    
    for equipment in supplier_info["equipment_list"]:
        row = table.add_row()
        row.cells[0].text = equipment["name"]
        row.cells[1].text = equipment["spec"]
        row.cells[2].text = str(equipment["quantity"])
        row.cells[3].text = equipment["manufacturer"]
    
    # 交付计划
    doc.add_heading('六、交付计划', 1)
    doc.add_paragraph('项目进度计划：')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = '阶段'
    header_cells[1].text = '工作内容'
    header_cells[2].text = '完成时间'
    
    for phase in supplier_info["delivery_plan"]:
        row = table.add_row()
        row.cells[0].text = phase["stage"]
        row.cells[1].text = phase["content"]
        row.cells[2].text = phase["time"]
    
    # 商务报价
    doc.add_heading('七、商务报价', 1)
    doc.add_paragraph('报价明细：')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = '项目'
    header_cells[1].text = '说明'
    header_cells[2].text = '金额（万元）'
    
    for item in supplier_info["price_breakdown"]:
        row = table.add_row()
        row.cells[0].text = item["item"]
        row.cells[1].text = item["description"]
        row.cells[2].text = str(item["amount"])
    
    row = table.add_row()
    row.cells[0].text = '合计'
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[1].text = ''
    row.cells[2].text = str(supplier_info["total_price"])
    row.cells[2].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph(f'付款方式：{supplier_info["payment_terms"]}')
    
    # 服务保障（不符合格式时缺少此章节）
    if format_compliant:
        doc.add_heading('八、服务保障', 1)
        doc.add_paragraph('1. 售后服务：')
        doc.add_paragraph(supplier_info["after_sales"], style='List Bullet')
        
        doc.add_paragraph('2. 技术支持：')
        doc.add_paragraph(supplier_info["technical_support"], style='List Bullet')
        
        doc.add_paragraph('3. 培训服务：')
        doc.add_paragraph(supplier_info["training"], style='List Bullet')
    # 不符合格式时，直接跳过"服务保障"章节
    
    # 附件清单（不符合格式时，章节编号应为"八"）
    attachment_heading = '九、附件清单' if format_compliant else '八、附件清单'
    doc.add_heading(attachment_heading, 1)
    attachments = supplier_info.get("attachments", [])
    if attachments:
        for attachment in attachments:
            doc.add_paragraph(f'• {attachment}', style='List Bullet')
    else:
        doc.add_paragraph('（无附件）')
    
    # 保存文件到招标书目录
    suffix = "_投标文件" if format_compliant else "_投标文件_格式不符合"
    output_dir = Path(__file__).parent / '招标书'
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / f'{supplier_name}{suffix}.docx'
    doc.save(str(output_path))
    status = "符合格式" if format_compliant else "不符合格式（缺少服务保障章节）"
    print(f'投标文件已生成：{output_path} ({status})')
    return output_path


def main():
    """主函数：生成三家供应商的投标文件"""
    print('开始生成投标文件...')
    
    # 供应商1：优秀供应商（技术方案好、价格合理、资质齐全、业绩好）
    supplier1 = {
        "total_price": 8200,  # 价格较低，有竞争力
        "delivery_period": 16,  # 交付周期合理
        "warranty_period": 36,
        "address": "陕西省西安市高新区航空路168号",
        "registered_capital": 10000,  # 注册资本较高
        "established_date": "2008年3月",
        "registration_number": "91610000123456789X",
        "legal_representative": "李建国",
        "company_intro": "我公司是一家专业从事航空结构件制造的高新技术企业，拥有16年的航空制造经验。公司专注于飞机机身、机翼等大型结构件的研发、制造和装配，已为多家航空企业提供优质产品和服务。公司拥有完善的质量管理体系和先进的生产设备，在航空制造领域享有良好声誉。",
        "certifications": [
            "AS9100D质量管理体系认证（2023年）",
            "NADCAP认证（热处理、无损检测、化学处理）",
            "ISO 14001环境管理体系认证",
            "ISO 45001职业健康安全管理体系认证",
            "FAA认证"
        ],
        "manufacturing_process": "采用国际先进的数字化智能制造技术，工艺流程如下：\n1. 材料预处理：建立完善的原材料入厂检验体系，对航空铝合金进行化学成分分析、力学性能检测和表面质量检验\n2. 蒙皮成形：采用五轴数控拉伸成形机，配备先进的回弹补偿算法，确保成形精度达到±0.1mm\n3. 机械加工：使用大型五轴加工中心，采用自适应加工技术，加工精度达到±0.02mm，表面粗糙度Ra≤0.8μm\n4. 装配：采用数字化装配系统，使用激光跟踪仪进行精确定位，装配精度达到±0.05mm\n5. 检测：使用高精度三坐标测量机、X射线检测设备、超声波相控阵检测设备进行全面检测，关键尺寸100%检测\n6. 表面处理：采用环保型阳极氧化和喷涂处理工艺，满足防腐和环保要求",
        "quality_system": "建立了完善的AS9100D质量管理体系，设立了独立的质量部门，配备专职质量工程师20人。建立了从原材料入厂到产品出厂的全过程数字化质量控制系统，实现质量数据的实时采集、分析和追溯。建立了供应商质量审核体系，关键工序设置SPC统计过程控制。",
        "testing_equipment": [
            "高精度三坐标测量机（精度0.0005mm，测量范围12m×8m×5m）",
            "X射线实时成像检测设备（最大穿透厚度80mm）",
            "超声波相控阵检测设备",
            "渗透检测设备",
            "工业CT检测设备",
            "材料拉伸试验机",
            "硬度计",
            "光谱分析仪"
        ],
        "control_points": [
            "建立完善的供应商质量审核体系，确保原材料质量符合标准",
            "关键尺寸100%检测，采用三坐标测量机进行全尺寸检测",
            "焊接部位100%无损检测，采用X射线和超声波双重检测",
            "装配精度实时监控，采用激光跟踪仪进行在线检测",
            "建立质量追溯系统，实现产品全生命周期追溯",
            "关键工序设置SPC统计过程控制，确保过程稳定"
        ],
        "team_size": 55,  # 团队规模较大
        "senior_engineers": 12,  # 高级工程师较多
        "engineers": 20,
        "technicians": 21,
        "project_manager": "王总工程师（教授级高工，25年航空制造经验，曾参与C919、ARJ21等多个国家重点型号项目）",
        "similar_projects": [
            {
                "name": "C919原型机机身段制造项目",
                "period": "2015-2017年",
                "content": "参与C919首架原型机机身中段制造，交付1架份，项目质量优良",
                "amount": 12000
            },
            {
                "name": "ARJ21机身段制造项目",
                "period": "2021-2023年",
                "content": "承担ARJ21支线客机机身中段制造，交付5架份，按期交付",
                "amount": 8500
            },
            {
                "name": "某型公务机机身制造项目",
                "period": "2019-2021年",
                "content": "大型公务机机身结构件制造，交付4架份，获得客户好评",
                "amount": 7200
            },
            {
                "name": "某型运输机机身段制造项目",
                "period": "2020-2023年",
                "content": "中型运输机机身段制造，交付6架份，质量稳定",
                "amount": 11200
            },
            {
                "name": "某型无人机机身制造项目",
                "period": "2022-2023年",
                "content": "大型无人机机身结构件制造，交付3套，技术先进",
                "amount": 4200
            }
        ],
        "equipment_list": [
            {"name": "五轴数控拉伸成形机", "spec": "工作台面8m×4m，精度±0.1mm", "quantity": 2, "manufacturer": "德国"},
            {"name": "五轴加工中心", "spec": "工作台面10m×5m，精度±0.01mm", "quantity": 4, "manufacturer": "瑞士"},
            {"name": "高精度三坐标测量机", "spec": "测量范围12m×8m×5m，精度0.0005mm", "quantity": 2, "manufacturer": "德国"},
            {"name": "X射线实时成像检测设备", "spec": "最大穿透厚度80mm", "quantity": 1, "manufacturer": "美国"},
            {"name": "数字化装配系统", "spec": "激光跟踪仪定位，精度±0.05mm", "quantity": 1, "manufacturer": "法国"},
            {"name": "工业CT检测设备", "spec": "检测范围直径1.5m", "quantity": 1, "manufacturer": "德国"}
        ],
        "delivery_plan": [
            {"stage": "第一阶段", "content": "技术准备、工艺验证、工装设计制造", "time": "合同签订后3个月"},
            {"stage": "第二阶段", "content": "首件试制、工艺优化、质量验证", "time": "第4-8个月"},
            {"stage": "第三阶段", "content": "小批量试制、质量检测、工艺固化", "time": "第9-13个月"},
            {"stage": "第四阶段", "content": "批量生产、总装、检测", "time": "第14-15个月"},
            {"stage": "第五阶段", "content": "最终检测、交付准备", "time": "第16个月"}
        ],
        "price_breakdown": [
            {"item": "材料费", "description": "航空级铝合金、钛合金、紧固件等", "amount": 3100},
            {"item": "加工费", "description": "成形、机械加工、装配", "amount": 3700},
            {"item": "检测费", "description": "质量检测、无损检测、试验", "amount": 700},
            {"item": "管理费", "description": "项目管理、质量管理、技术支持", "amount": 400},
            {"item": "利润", "description": "", "amount": 300}
        ],
        "payment_terms": "合同签订后支付25%预付款，首件验收合格后支付35%，批量交付验收合格后支付35%，质保期满后支付5%",  # 付款方式对招标方有利
        "after_sales": "提供36个月质保服务，质保期内免费维修和更换。设立专门的服务团队，24小时响应，48小时内到达现场。建立完善的备件库，确保及时供应。",
        "technical_support": "提供全方位技术支持，包括工艺咨询、技术培训、现场指导、问题解决等。配备3名高级技术支持工程师，其中1名常驻现场。建立技术专家库，可随时提供远程技术支持。",
        "training": "提供系统化培训，包括技术培训、操作培训、维护培训、质量管理培训等，总培训时间不少于60小时，并提供培训资料和证书。可根据需求定制培训内容。",
        "attachments": [
            "营业执照复印件",
            "AS9100认证证书",
            "NADCAP认证证书",
            "FAA认证证书",
            "类似业绩合同复印件（5份）",
            "主要设备清单及证明",
            "技术团队简历",
            "质量体系文件"
        ]
    }
    print('\n=== 生成供应商1投标文件（优秀供应商） ===')
    create_supplier_bid_document("西安航空制造有限公司", supplier1, format_compliant=True)
    
    # 供应商2：中等供应商（技术方案一般、价格稍高、资质一般、业绩较少）
    supplier2 = {
        "total_price": 9100,  # 价格较高
        "delivery_period": 17,  # 交付周期稍长
        "warranty_period": 36,
        "address": "江苏省苏州市工业园区航空科技园88号",
        "registered_capital": 7500,  # 注册资本中等
        "established_date": "2012年6月",
        "registration_number": "91320500123456789Y",
        "legal_representative": "张明华",
        "company_intro": "我公司是专业从事航空精密制造的企业，在航空结构件制造领域具有12年经验。公司拥有一定的制造设备和检测设备，已通过相关认证。公司注重产品质量，致力于为客户提供优质服务。",
        "certifications": [
            "AS9100D质量管理体系认证（2022年）",
            "ISO 9001质量管理体系认证",
            "ISO 14001环境管理体系认证"
        ],
        "manufacturing_process": "采用成熟的制造工艺：\n1. 材料预处理：对航空铝合金进行检验\n2. 蒙皮成形：采用数控拉伸成形机进行成形\n3. 机械加工：使用五轴加工中心进行加工，加工精度±0.05mm\n4. 装配：采用数字化装配系统进行装配\n5. 检测：使用三坐标测量机、X射线检测设备进行检测\n6. 表面处理：进行阳极氧化和喷涂处理",
        "quality_system": "建立了AS9100质量管理体系，质量部门配备12名专业人员。建立了基本的质量控制流程，对关键工序进行检验。",
        "testing_equipment": [
            "三坐标测量机（精度0.001mm）",
            "X射线检测设备",
            "超声波检测设备",
            "渗透检测设备",
            "材料试验机"
        ],
        "control_points": [
            "原材料入厂检验",
            "关键尺寸检测",
            "焊接部位无损检测",
            "装配精度检测",
            "最终产品检测"
        ],
        "team_size": 35,  # 团队规模较小
        "senior_engineers": 5,  # 高级工程师较少
        "engineers": 12,
        "technicians": 16,
        "project_manager": "陈总工程师（高级工程师，15年航空制造经验）",
        "similar_projects": [
            {
                "name": "某型无人机机身制造项目",
                "period": "2021-2022年",
                "content": "中型无人机机身结构件制造，交付5套",
                "amount": 2800
            },
            {
                "name": "某型直升机机身段制造项目",
                "period": "2022-2023年",
                "content": "轻型直升机机身段制造，交付8架份",
                "amount": 3600
            }
        ],
        "equipment_list": [
            {"name": "五轴加工中心", "spec": "工作台面8m×4m，精度±0.05mm", "quantity": 2, "manufacturer": "国产"},
            {"name": "数控拉伸成形机", "spec": "工作台面6m×3m", "quantity": 1, "manufacturer": "国产"},
            {"name": "三坐标测量机", "spec": "测量范围10m×6m×4m，精度0.001mm", "quantity": 1, "manufacturer": "国产"},
            {"name": "X射线检测设备", "spec": "最大穿透厚度50mm", "quantity": 1, "manufacturer": "国产"}
        ],
        "delivery_plan": [
            {"stage": "第一阶段", "content": "技术准备、工装设计制造", "time": "合同签订后4个月"},
            {"stage": "第二阶段", "content": "首件试制、工艺验证", "time": "第5-9个月"},
            {"stage": "第三阶段", "content": "批量生产、质量检测", "time": "第10-14个月"},
            {"stage": "第四阶段", "content": "总装、检测、交付", "time": "第15-17个月"}
        ],
        "price_breakdown": [
            {"item": "材料费", "description": "航空铝合金、紧固件等", "amount": 3500},
            {"item": "加工费", "description": "成形、机械加工、装配", "amount": 4000},
            {"item": "检测费", "description": "质量检测、无损检测", "amount": 700},
            {"item": "管理费", "description": "项目管理、质量管理", "amount": 500},
            {"item": "利润", "description": "", "amount": 400}
        ],
        "payment_terms": "合同签订后支付30%预付款，首件验收合格后支付30%，批量交付验收合格后支付35%，质保期满后支付5%",  # 付款方式一般
        "after_sales": "提供36个月质保服务，质保期内免费维修。设立服务热线，接到通知后72小时内响应。",
        "technical_support": "提供技术支持服务，包括工艺指导、技术培训等。配备1名技术支持工程师，根据需要进行现场支持。",
        "training": "提供操作培训和维护培训，培训时间不少于30小时。",
        "attachments": [
            "营业执照复印件",
            "AS9100认证证书",
            "类似业绩合同复印件（2份）",
            "主要设备清单",
            "技术团队简历",
            "质量体系文件"
        ]
    }
    print('\n=== 生成供应商2投标文件（中等供应商） ===')
    create_supplier_bid_document("苏州精密航空科技有限公司", supplier2, format_compliant=True)
    
    # 供应商3：格式不符合（缺少服务保障章节）
    supplier3 = {
        "total_price": 8800,  # 价格中等
        "delivery_period": 17,
        "warranty_period": 36,
        "address": "四川省成都市双流区航空产业园99号",
        "registered_capital": 6500,
        "established_date": "2013年8月",
        "registration_number": "91510100123456789Z",
        "legal_representative": "刘强",
        "company_intro": "我公司是专业从事航空结构件制造的现代化企业，在航空制造领域深耕11年。公司拥有完整的航空制造产业链，从原材料加工到最终装配，具备一站式服务能力。公司注重技术创新和人才培养，与多所高校建立了产学研合作关系。",
        "certifications": [
            "AS9100D质量管理体系认证（2023年）",
            "ISO 9001质量管理体系认证",
            "ISO 14001环境管理体系认证"
        ],
        "manufacturing_process": "采用成熟的制造工艺：\n1. 材料准备：建立材料库，对原材料进行分类管理和检验\n2. 下料：采用激光切割和数控下料，提高材料利用率\n3. 成形：采用滚弯成形和拉伸成形相结合的方式\n4. 机械加工：采用国产和进口设备相结合，确保加工质量\n5. 装配：采用传统装配工艺结合数字化辅助，确保装配精度\n6. 检测：建立完善的检测体系，确保产品质量\n7. 表面处理：采用成熟的表面处理工艺",
        "quality_system": "建立了符合AS9100标准的质量管理体系，质量部门配备15名专业人员。建立了质量责任制，每个工序都有明确的质量责任人。定期开展质量培训和审核活动。",
        "testing_equipment": [
            "三坐标测量机（精度0.002mm）",
            "X射线检测设备",
            "超声波检测设备",
            "渗透检测设备",
            "材料试验机",
            "硬度计"
        ],
        "control_points": [
            "建立原材料供应商审核制度",
            "关键工序设置检验点，实行首件检验",
            "关键尺寸采用三坐标测量机检测",
            "建立质量档案，实现质量追溯",
            "定期进行质量分析和改进"
        ],
        "team_size": 32,
        "senior_engineers": 5,
        "engineers": 11,
        "technicians": 14,
        "project_manager": "赵总工程师（高级工程师，16年航空制造经验）",
        "similar_projects": [
            {
                "name": "某型教练机机身段制造项目",
                "period": "2020-2022年",
                "content": "教练机机身中段制造，交付8架份",
                "amount": 4800
            },
            {
                "name": "某型无人机机身制造项目",
                "period": "2021-2023年",
                "content": "中型无人机机身结构件制造，交付6套",
                "amount": 3200
            }
        ],
        "equipment_list": [
            {"name": "五轴加工中心", "spec": "工作台面6m×3m", "quantity": 2, "manufacturer": "国产"},
            {"name": "数控滚弯机", "spec": "最大滚弯长度8m", "quantity": 2, "manufacturer": "国产"},
            {"name": "拉伸成形机", "spec": "工作台面5m×3m", "quantity": 1, "manufacturer": "国产"},
            {"name": "三坐标测量机", "spec": "测量范围8m×5m×3m", "quantity": 1, "manufacturer": "国产"},
            {"name": "X射线检测设备", "spec": "最大穿透厚度40mm", "quantity": 1, "manufacturer": "国产"}
        ],
        "delivery_plan": [
            {"stage": "第一阶段", "content": "技术准备、工装设计制造", "time": "合同签订后4个月"},
            {"stage": "第二阶段", "content": "首件试制、工艺验证", "time": "第5-9个月"},
            {"stage": "第三阶段", "content": "批量生产、质量检测", "time": "第10-14个月"},
            {"stage": "第四阶段", "content": "总装、检测、交付", "time": "第15-17个月"}
        ],
        "price_breakdown": [
            {"item": "材料费", "description": "航空铝合金、紧固件等", "amount": 3300},
            {"item": "加工费", "description": "成形、机械加工、装配", "amount": 3800},
            {"item": "检测费", "description": "质量检测、无损检测", "amount": 600},
            {"item": "管理费", "description": "项目管理、质量管理", "amount": 300},
            {"item": "利润", "description": "", "amount": 800}
        ],
        "payment_terms": "合同签订后支付30%预付款，首件验收合格后支付30%，批量交付验收合格后支付35%，质保期满后支付5%",
        "after_sales": "提供36个月质保服务，质保期内免费维修。设立服务热线，接到通知后72小时内响应。",
        "technical_support": "提供技术支持服务，包括工艺指导、技术培训等。配备1名技术支持工程师，根据需要进行现场支持。",
        "training": "提供必要的操作培训和维护培训，培训时间不少于30小时。",
        "attachments": [
            "营业执照复印件",
            "AS9100认证证书",
            "类似业绩合同复印件（2份）",
            "主要设备清单",
            "技术团队简历",
            "质量体系文件"
        ]
    }
    print('\n=== 生成供应商3投标文件（格式不符合，缺少服务保障章节） ===')
    create_supplier_bid_document("成都航空结构件制造股份有限公司", supplier3, format_compliant=False)
    
    print('\n所有投标文件生成完成！')


if __name__ == '__main__':
    main()

