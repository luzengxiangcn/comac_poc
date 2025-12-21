"""
使用 DeepSeek 模型生成招标文件
基于技术需求文件和采购部门标书要求，使用 LLM 生成 markdown 格式的标书，然后转换为 docx
"""
import asyncio
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 添加项目根目录到路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from src.comac_purchase.model_session.model import openai_client
from src.comac_purchase.config import settings
from openai import AsyncOpenAI


def read_docx_to_text(file_path: Path) -> str:
    """读取 docx 文件内容并转换为纯文本（保留结构）"""
    from docx import Document
    
    doc = Document(file_path)
    text_lines = []
    
    for element in doc.element.body:
        # 处理段落
        if element.tag.endswith('p'):
            from docx.oxml.text.paragraph import CT_P
            from docx.text.paragraph import Paragraph
            
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                text = para.text.strip()
                if text:
                    # 检查是否是标题
                    if para.style and ('Heading' in para.style.name or '标题' in para.style.name):
                        level = _get_heading_level(para.style.name)
                        if level > 0:
                            text_lines.append('#' * level + ' ' + text)
                        else:
                            text_lines.append(text)
                    else:
                        text_lines.append(text)
        
        # 处理表格
        elif element.tag.endswith('tbl'):
            from docx.oxml.table import CT_Tbl
            from docx.table import Table
            
            if isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table_text = _table_to_text(table)
                if table_text:
                    text_lines.append('')
                    text_lines.append(table_text)
                    text_lines.append('')
    
    return '\n'.join(text_lines)


def _get_heading_level(style_name: str) -> int:
    """从样式名称获取标题级别"""
    if 'Heading 1' in style_name or '标题 1' in style_name or 'Title' in style_name:
        return 1
    elif 'Heading 2' in style_name or '标题 2' in style_name:
        return 2
    elif 'Heading 3' in style_name or '标题 3' in style_name:
        return 3
    elif 'Heading 4' in style_name or '标题 4' in style_name:
        return 4
    elif 'Heading 5' in style_name or '标题 5' in style_name:
        return 5
    elif 'Heading 6' in style_name or '标题 6' in style_name:
        return 6
    return 0


def _table_to_text(table) -> str:
    """将表格转换为文本格式"""
    text_lines = []
    for row in table.rows:
        row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        text_lines.append(' | '.join(row_data))
    return '\n'.join(text_lines)


def read_markdown_file(file_path: Path) -> str:
    """读取 markdown 文件内容"""
    return file_path.read_text(encoding='utf-8')


async def generate_tender_with_llm(
    technical_requirement_path: Path,
    procurement_requirement_path: Path
) -> str:
    """使用 LLM 生成招标文件（markdown 格式）
    
    Args:
        technical_requirement_path: 技术需求文件路径
        procurement_requirement_path: 采购部门标书要求文件路径
        
    Returns:
        markdown 格式的招标文件内容
    """
    # 读取需求文件内容
    print("正在读取需求文件...")
    
    if technical_requirement_path.suffix == '.docx':
        technical_content = read_docx_to_text(technical_requirement_path)
    elif technical_requirement_path.suffix == '.md':
        technical_content = read_markdown_file(technical_requirement_path)
    else:
        raise ValueError(f"不支持的文件格式: {technical_requirement_path.suffix}")
    
    if procurement_requirement_path.suffix == '.docx':
        procurement_content = read_docx_to_text(procurement_requirement_path)
    elif procurement_requirement_path.suffix == '.md':
        procurement_content = read_markdown_file(procurement_requirement_path)
    else:
        raise ValueError(f"不支持的文件格式: {procurement_requirement_path.suffix}")
    
    # 构建 prompt
    prompt = f"""你是一位专业的招标文件编写专家。请根据以下技术部门的需求文件和采购部门的标书要求，生成一份完整的、专业的招标文件。

## 技术部门需求文件

{technical_content}

## 采购部门标书要求

{procurement_content}

## 任务要求

请根据以上两个文件的内容，生成一份完整的招标文件（Markdown格式），因为后续会转换为docx格式，所以请不要包含任何格式标记。招标文件应包含以下章节结构：

1. **标题**：根据技术需求文件中的项目名称生成，格式为"[项目名称]招标文件"（居中显示）
2. **一、项目基本信息**：从技术需求文件中提取项目名称、项目编号等信息，从采购部门要求中提取招标单位、招标日期、投标截止时间等信息
3. **二、项目概述**：完整引用技术需求文件中的项目概述内容
4. **三、技术要求**：完整引用技术需求文件中的技术要求，包括所有子章节（如材料要求、制造工艺要求、质量要求、交付要求等）
5. **四、商务要求**：完整引用采购部门要求中的商务要求内容
6. **五、投标人资格要求**：完整引用采购部门要求中的投标人资格要求内容
7. **六、评分细则**：完整引用采购部门要求中的评分细则，并以表格形式呈现评分标准
8. **七、标书格式要求**：完整引用采购部门要求中的标书格式要求内容
9. **八、其他说明**：完整引用采购部门要求中的其他说明内容
10. **九、联系方式**：完整引用采购部门要求中的联系方式内容

## 注意事项

1. 保持专业、严谨的文档风格
2. 确保所有章节内容完整、逻辑清晰
3. 评分细则需要以表格形式呈现，表格应包含：评分项目、分值、评分标准、备注等列
4. 标书格式要求需要详细说明章节结构和格式评分标准
5. 使用标准的 Markdown 格式，包括标题（#）、列表（- 或 1.）、表格（|）等
6. 确保所有信息准确，不要遗漏任何重要内容
7. 项目名称、项目编号等具体信息应从技术需求文件中提取，不要自行编造
8. 招标单位、联系方式等信息应从采购部门要求中提取

请直接输出完整的招标文件 Markdown 内容，不要包含任何额外的说明或注释。"""

    print("正在调用 DeepSeek 模型生成招标文件...")
    print("=" * 60)
    print("开始流式生成内容（实时输出）：")
    print("=" * 60)
    print()
    
    # 创建原始客户端用于流式输出（避免重试机制对流的干扰）
    # 如果需要重试，可以在外层包装
    raw_client = AsyncOpenAI(
        api_key=settings.LLM_API_KEY,
        base_url=settings.LLM_BASE_URL
    )
    
    # 调用 LLM（流式模式）
    try:
        stream = await raw_client.chat.completions.create(
            model="deepseek-ai/DeepSeek-V3.2-Exp",
            messages=[
                {
                    "role": "system",
                    "content": "你是一位专业的招标文件编写专家，擅长根据技术需求和采购要求编写完整、专业的招标文件。"
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.3,
            max_tokens=8000,
            stream=True  # 启用流式输出
        )
        
        # 收集完整内容
        markdown_content = ""
        
        # 处理流式响应
        async for chunk in stream:
            if chunk.choices and len(chunk.choices) > 0:
                delta = chunk.choices[0].delta
                if hasattr(delta, 'content') and delta.content:
                    content = delta.content
                    markdown_content += content
                    # 实时输出内容（不换行，实现打字机效果）
                    print(content, end='', flush=True)
        
        print()
        print()
        print("=" * 60)
        print("招标文件生成完成！")
        print("=" * 60)
        
        return markdown_content
    
    except Exception as e:
        print()
        print(f"\n流式生成出错，尝试使用非流式模式...")
        # 如果流式失败，回退到非流式模式
        response = await openai_client.chat.completions.create(
            model="deepseek-ai/DeepSeek-V3.2-Exp",
            messages=[
                {
                    "role": "system",
                    "content": "你是一位专业的招标文件编写专家，擅长根据技术需求和采购要求编写完整、专业的招标文件。"
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.3,
            max_tokens=8000
        )
        
        markdown_content = response.choices[0].message.content
        print("招标文件生成完成！")
        return markdown_content


def markdown_to_docx(markdown_content: str, output_path: Path):
    """将 Markdown 内容转换为 docx 文件
    
    Args:
        markdown_content: Markdown 格式的内容
        output_path: 输出文件路径
    """
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 设置标题样式
    for i in range(1, 7):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = '黑体'
        heading_font.size = Pt(18 - i * 2)
        heading_font.bold = True
    
    lines = markdown_content.split('\n')
    i = 0
    current_list_style = None  # 当前列表样式
    
    def add_formatted_text(paragraph, text: str):
        """向段落添加格式化的文本（支持粗体）"""
        # 处理粗体 **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part:
                paragraph.add_run(part)
    
    while i < len(lines):
        line = lines[i]
        stripped_line = line.strip()
        
        if not stripped_line:
            i += 1
            continue
        
        # 处理标题
        if stripped_line.startswith('#'):
            level = 0
            while level < len(stripped_line) and stripped_line[level] == '#':
                level += 1
            title_text = stripped_line[level:].strip()
            
            # 移除标题中的粗体标记（标题本身就是粗体）
            title_text = re.sub(r'\*\*(.*?)\*\*', r'\1', title_text)
            
            if level == 1:
                heading = doc.add_heading(title_text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(title_text, min(level - 1, 6))
            current_list_style = None
        
        # 处理表格
        elif stripped_line.startswith('|') and '|' in stripped_line[1:]:
            # 收集表格行
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                # 跳过分隔行（如 |---|---|）
                if not re.match(r'^\|[\s\-:]+\|', row_line):
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    # 移除单元格中的 markdown 格式标记
                    cells = [re.sub(r'\*\*(.*?)\*\*', r'\1', cell) for cell in cells]
                    table_rows.append(cells)
                i += 1
            i -= 1  # 回退一行
            
            # 创建表格
            if table_rows:
                num_cols = max(len(row) for row in table_rows) if table_rows else 1
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = 'Light Grid Accent 1'
                
                # 添加表头
                if table_rows:
                    header_cells = table.rows[0].cells
                    for j in range(num_cols):
                        if j < len(table_rows[0]):
                            header_cells[j].text = table_rows[0][j]
                        else:
                            header_cells[j].text = ''
                    
                    # 添加数据行
                    for row_data in table_rows[1:]:
                        row = table.add_row()
                        for j in range(num_cols):
                            if j < len(row_data):
                                row.cells[j].text = row_data[j]
                            else:
                                row.cells[j].text = ''
            current_list_style = None
        
        # 处理无序列表
        elif re.match(r'^[\-\*\+]\s+', stripped_line):
            list_text = re.sub(r'^[\-\*\+]\s+', '', stripped_line)
            p = doc.add_paragraph(list_text, style='List Bullet')
            # 处理列表项中的粗体
            if '**' in list_text:
                p.clear()
                add_formatted_text(p, list_text)
            current_list_style = 'bullet'
        
        # 处理有序列表
        elif re.match(r'^\d+\.\s+', stripped_line):
            list_text = re.sub(r'^\d+\.\s+', '', stripped_line)
            p = doc.add_paragraph(list_text, style='List Number')
            # 处理列表项中的粗体
            if '**' in list_text:
                p.clear()
                add_formatted_text(p, list_text)
            current_list_style = 'number'
        
        # 处理普通段落
        else:
            # 检查是否是列表的延续（缩进文本）
            if line.startswith('   ') or line.startswith('\t'):
                # 作为前一个列表项的延续
                text = stripped_line
                p = doc.add_paragraph(text)
                if '**' in text:
                    p.clear()
                    add_formatted_text(p, text)
            else:
                p = doc.add_paragraph()
                add_formatted_text(p, stripped_line)
            current_list_style = None
        
        i += 1
    
    # 保存文件
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f'招标文件已保存为 docx 格式：{output_path}')


async def main():
    """主函数"""
    # 需求文件路径
    requirement_dir = Path(__file__).parent / '需求文件'
    technical_requirement_path = requirement_dir / '技术需求文件_C919机身段制造项目.docx'
    procurement_requirement_path = requirement_dir / '采购部门标书要求.md'
    
    # 检查文件是否存在
    if not technical_requirement_path.exists():
        print(f"错误：技术需求文件不存在：{technical_requirement_path}")
        return
    
    if not procurement_requirement_path.exists():
        # 尝试查找 docx 格式
        procurement_requirement_path = requirement_dir / '采购部门标书要求.docx'
        if not procurement_requirement_path.exists():
            print(f"错误：采购部门标书要求文件不存在")
            return
    
    print("=" * 60)
    print("使用 DeepSeek 模型生成招标文件")
    print("=" * 60)
    
    try:
        # 生成 markdown 格式的招标文件
        markdown_content = await generate_tender_with_llm(
            technical_requirement_path,
            procurement_requirement_path
        )
        
        # 保存 markdown 文件
        markdown_output_path = Path(__file__).parent / 'LLM生成的招标文件.md'
        markdown_output_path.write_text(markdown_content, encoding='utf-8')
        print(f'Markdown 格式的招标文件已保存：{markdown_output_path}')
        
        # 转换为 docx
        docx_output_path = Path(__file__).parent / 'LLM生成的招标文件.docx'
        markdown_to_docx(markdown_content, docx_output_path)
        
        print("\n" + "=" * 60)
        print("招标文件生成完成！")
        print("=" * 60)
        print(f"Markdown 文件：{markdown_output_path}")
        print(f"DOCX 文件：{docx_output_path}")
        
    except Exception as e:
        print(f"生成招标文件时出错：{str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == '__main__':
    asyncio.run(main())

